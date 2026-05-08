import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime, timedelta

# 모델 라이브러리
from statsmodels.tsa.holtwinters import ExponentialSmoothing
from prophet import Prophet
from statsmodels.tsa.statespace.sarimax import SARIMAX
from sklearn.ensemble import RandomForestRegressor
from xgboost import XGBRegressor

# 보고서 생성
from docx import Document
from docx.shared import Inches
import io

# --- 1. 가상 데이터 생성 (테스트용) ---
def get_sample_data():
    dates = pd.date_range(start='2020-01-01', periods=48, freq='MS')
    data = {
        '날짜': dates,
        '수출금액': [100 + i + (i%12)*5 + np.random.randn()*5 for i in range(48)],
        '수입금액': [90 + i*0.8 + (i%12)*4 + np.random.randn()*4 for i in range(48)],
        '수출건수': [500 + i*10 for i in range(48)],
        '수입건수': [450 + i*8 for i in range(48)]
    }
    return pd.DataFrame(data)

# --- 2. 수요예측 모델 엔진 ---
def forecast_logic(df, column, model_type, p_period, params):
    series = df[column].values
    last_date = df['날짜'].max()
    future_dates = pd.date_range(start=last_date + pd.DateOffset(months=1), periods=p_period, freq='MS')
    
    forecast_val = []
    lower_bound = []
    upper_bound = []

    if model_type == "Exponential Smoothing":
        model = ExponentialSmoothing(series, trend='add', seasonal='add', seasonal_periods=12).fit()
        forecast_val = model.forecast(p_period)
        lower_bound = forecast_val * 0.95 # 단순화된 신뢰구간
        upper_bound = forecast_val * 1.05

    elif model_type == "Facebook Prophet":
        pdf = df.rename(columns={'날짜': 'ds', column: 'y'})
        m = Prophet(yearly_seasonality=True, weekly_seasonality=False, daily_seasonality=False)
        m.fit(pdf)
        future = m.make_future_dataframe(periods=p_period, freq='MS')
        forecast = m.predict(future)
        forecast_val = forecast['yhat'].iloc[-p_period:].values
        lower_bound = forecast['yhat_lower'].iloc[-p_period:].values
        upper_bound = forecast['yhat_upper'].iloc[-p_period:].values

    elif model_type == "XGBoost":
        # 시차 변수(Lag) 생성 간단 모델링
        y = series
        X = np.arange(len(y)).reshape(-1, 1)
        model = XGBRegressor().fit(X, y)
        future_X = np.arange(len(y), len(y) + p_period).reshape(-1, 1)
        forecast_val = model.predict(future_X)
        lower_bound = forecast_val * 0.92
        upper_bound = forecast_val * 1.08

    return future_dates, forecast_val, lower_bound, upper_bound

# --- 3. UI 구성 ---
st.set_page_config(page_title="E-commerce Forecast", layout="wide")

# 모델 설명 커스텀 HTML/CSS
st.markdown("""
    <style>
    .desc-box { background-color: #f0f2f6; padding: 10px; border-radius: 10px; margin-bottom: 20px; }
    .dark .desc-box { background-color: #262730; }
    </style>
""", unsafe_allow_html=True)

st.title("🌐 국가별 전자상거래 수출입 수요예측 시스템")

with st.sidebar:
    st.header("📂 데이터 및 설정")
    file = st.file_uploader("엑셀 파일 업로드 (날짜, 수출금액, 수입금액 등 포함)", type=['xlsx'])
    
    st.subheader("⏱️ 기간 설정")
    obs_range = st.slider("분석 관측 기간", 6, 60, 36)
    pred_range = st.selectbox("예측 기간", [3, 6, 12], index=2)
    
    st.subheader("🤖 모델 선택")
    selected_models = st.multiselect(
        "사용할 모델 (복수 선택 가능)", 
        ["Exponential Smoothing", "Facebook Prophet", "SARIMA", "XGBoost", "Random Forest"],
        default=["Facebook Prophet"]
    )

# 메인 분석 로직
data = get_sample_data() if file is None else pd.read_excel(file)

# 그래프 그리기
fig = go.Figure()

# 실제 데이터 추가
fig.add_trace(go.Scatter(x=data['날짜'], y=data['수출금액'], name="실제 수출액 ($)", line=dict(color='blue', width=3)))
fig.add_trace(go.Scatter(x=data['날짜'], y=data['수입금액'], name="실제 수입액 ($)", line=dict(color='red', width=3)))

# 예측 수행 및 그래프 추가
for m_name in selected_models:
    try:
        f_date, f_val, low, high = forecast_logic(data, '수출금액', m_name, pred_range, {})
        
        # 예측 선
        fig.add_trace(go.Scatter(x=f_date, y=f_val, name=f"{m_name} 예측", line=dict(dash='dash')))
        
        # 신뢰 구간 (선택 사항)
        fig.add_trace(go.Scatter(
            x=np.concatenate([f_date, f_date[::-1]]),
            y=np.concatenate([high, low[::-1]]),
            fill='toself',
            fillcolor='rgba(0,100,80,0.1)',
            line=dict(color='rgba(255,255,255,0)'),
            hoverinfo="skip",
            showlegend=False,
            name="예상 범위"
        ))
    except:
        st.warning(f"{m_name} 모델 계산 중 오류가 발생했습니다.")

fig.update_layout(
    hovermode="x unified",
    xaxis_title="연월",
    yaxis_title="금액 (USD $)",
    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
)

st.plotly_chart(fig, use_container_width=True)

# 보고서 생성 버튼
if st.button("📝 Word 보고서 생성 및 다운로드"):
    doc = Document()
    doc.add_heading('수요예측 분석 보고서', 0)
    doc.add_paragraph(f'분석 일시: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_section()
    doc.add_paragraph("선택된 모델을 기반으로 한 향후 12개월 수출입 예측 결과입니다...")
    
    # 메모리에 저장 후 다운로드
    bio = io.BytesIO()
    doc.save(bio)
    st.download_button(
        label="Word 파일 받기",
        data=bio.getvalue(),
        file_name=f"Forecast_Report_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )